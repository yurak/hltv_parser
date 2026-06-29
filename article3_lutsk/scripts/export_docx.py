"""
export_docx.py — render paper/manuscript_revised.md to paper/manuscript.docx
with journal №85 (ЛНТУ) formatting:
  A4; mirrored margins top/bottom 1.5, left 2.5, right 2 cm; Times New Roman;
  body 11 pt (indent 1.25 cm), abstracts 9 pt (indent 1 cm); single spacing; no page numbers.
Embeds the four figures from outputs/figures/ at their callouts.

Tailored Markdown subset parser (not general): handles УДК, author lines, #/##/### headings,
paragraphs, the markdown table, blockquotes, separators, inline **bold** and `code`.
"""
import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "paper", "manuscript_revised.md")
OUT = os.path.join(ROOT, "paper", "manuscript.docx")
FIGS = os.path.join(ROOT, "outputs", "figures")
FONT = "Times New Roman"

FIG_MAP = {
    "Рис. 1.": "invariance_summary.png",
    "Рис. 2.": "pca_projection.png",
    "Рис. 3.": "side_comparison.png",
    "Рис. 4.": "feature_distributions.png",
}

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)

def add_runs(p, text, size=11):
    """Parse inline **bold** and `code`; emit runs."""
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            set_font(p.add_run(tok[2:-2]), size=size, bold=True)
        elif tok.startswith("`") and tok.endswith("`"):
            set_font(p.add_run(tok[1:-1]), size=size)
        else:
            set_font(p.add_run(tok), size=size)

def para(doc, text, size=11, indent=1.25, align=None, bold=False, italic=False, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(indent) if indent else None
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align:
        p.alignment = align
    if bold or italic:
        set_font(p.add_run(text), size=size, bold=bold, italic=italic)
    else:
        add_runs(p, text, size=size)
    return p

def heading(doc, text, size=11, before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(text), size=size, bold=True)
    return p

def add_table(doc, rows):
    cols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.paragraphs[0].clear()
            set_font(c.paragraphs[0].add_run(cell.strip()), size=10, bold=(i == 0))
    return t

def embed_figure(doc, fname):
    path = os.path.join(FIGS, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def configure_page(doc):
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)         # A4
    s.top_margin, s.bottom_margin = Cm(1.5), Cm(1.5)
    s.left_margin, s.right_margin = Cm(2.5), Cm(2.0)
    s.header_distance, s.footer_distance = Cm(1.25), Cm(1.25)
    # mirrored margins
    pg = s._sectPr
    mirror = OxmlElement("w:mirrorMargins"); pg.insert(0, mirror)

def main():
    lines = open(SRC, encoding="utf-8").read().split("\n")
    doc = Document()
    # base style
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(11)
    configure_page(doc)

    in_abstract = False
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        stripped = ln.strip()

        if stripped == "" or stripped == "---":
            i += 1; continue

        # markdown table block
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip()); i += 1
            rows = []
            for r in block:
                if re.match(r"^\|[\s:|-]+\|$", r):  # separator row
                    continue
                cells = [c for c in r.strip().strip("|").split("|")]
                rows.append(cells)
            if rows:
                add_table(doc, rows)
            continue

        # headings
        if ln.startswith("# "):
            heading(doc, ln[2:].strip(), size=12, before=10)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue
        if ln.startswith("## "):
            title = ln[3:].strip()
            in_abstract = title.startswith(("Анотація", "Abstract"))
            heading(doc, title, size=11, before=8); i += 1; continue
        if ln.startswith("### "):
            heading(doc, ln[4:].strip(), size=11, before=6); i += 1; continue

        # blockquote (provenance / notes) — small italic
        if stripped.startswith(">"):
            txt = re.sub(r"^>\s?", "", stripped)
            if txt:
                para(doc, txt, size=9, indent=0, italic=False, space_after=1)
            i += 1; continue

        # УДК
        if stripped.startswith("УДК"):
            para(doc, stripped, size=11, indent=0, bold=True); i += 1; continue

        # author lines (start with **)
        if stripped.startswith("**") and ("ORCID" in stripped or "канд." in stripped or "аспірант" in stripped):
            para(doc, stripped, size=11, indent=0); i += 1; continue

        # figure callout — caption then embed
        m = re.match(r"\*\*(Рис\. \d\.)\*\*", stripped)
        if m and m.group(1) in FIG_MAP:
            embed_figure(doc, FIG_MAP[m.group(1)])
            para(doc, stripped, size=9, indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
            i += 1; continue

        # keywords line
        if stripped.startswith("**Ключові слова") or stripped.startswith("**Keywords"):
            para(doc, stripped, size=9, indent=0); i += 1; continue

        # regular paragraph
        size = 9 if in_abstract else 11
        indent = 1.0 if in_abstract else 1.25
        para(doc, stripped, size=size, indent=indent)
        i += 1

    doc.save(OUT)
    print("Saved", OUT)

if __name__ == "__main__":
    main()
